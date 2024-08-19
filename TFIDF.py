import os
import math
import thulac
import pandas as pd
from collections import defaultdict
import docx
import pysrt

# Global variables for input/output files, documents folder, and known words file
INPUT_FILENAME = 'Sample_Input_Media_To_Process.srt'
OUTPUT_FILENAME = 'Sample_Output.xlsx'
DOCUMENTS_FOLDER = 'Sample_Corpus'
KNOWN_WORDS_FILENAME = 'Sample_Known_Words.txt'

""" 
Loads the file that contains the words that the user already knows. These words will be filtered out of the TFIDF analysis.
"""
def load_known_words(filename):
    with open(filename, 'r', encoding='utf-8') as file:
        known_words = {line.strip() for line in file}
    return known_words

"""
Calculates the raw frequency of each word in the target media, and returns a dictionary with the count of each word.
"""
def calculate_target_media_word_frequency(filename):
    thu = thulac.thulac(seg_only=True)
    text = ""
    if filename.endswith('.txt'):
        with open(filename, 'r', encoding='utf-8') as file:
            text = file.read()
    elif filename.endswith('.srt'):
        subs = pysrt.open(filename, encoding='utf-8')
        text = ' '.join([sub.text for sub in subs])
    elif filename.endswith('.docx'):
        doc = docx.Document(filename)
        text = ' '.join([para.text for para in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format")
    
    words = thu.cut(text, text=True).split()
    word_freq = defaultdict(int)
    for word in words:
        word_freq[word] += 1
    return word_freq

"""
Goes through the folder containing all of the other reference documents that will be used in the comparison against the target media, and returns those documents as a list.
"""
def pull_reference_documents(folder_path):
    document_list = []
    for filename in os.listdir(folder_path):
        if filename.endswith(('.txt', '.srt', '.docx')):
            file_path = os.path.join(folder_path, filename)
            text = ""
            if filename.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            elif filename.endswith('.srt'):
                subs = pysrt.open(file_path, encoding='utf-8')
                text = ' '.join([sub.text for sub in subs])
            elif filename.endswith('.docx'):
                doc = docx.Document(file_path)
                text = ' '.join([para.text for para in doc.paragraphs])
            document_list.append({'DocumentName': filename, 'DocumentText': text})
    return document_list


"""
Calculates TFIDF for each word, as well as other intermediate statistics used in the TFIDF calculation.
"""
def calculate_statistics(word_list, document_list):
    total_num_documents = len(document_list)
    
    for word_entry in word_list:
        word = word_entry['word']
        num_documents_appear_in = sum(1 for doc in document_list if word in doc['DocumentText'])
        word_entry['# of Other Media Documents Appeared In'] = num_documents_appear_in
        word_entry['TFIDF LN Factor'] = math.log(total_num_documents / (num_documents_appear_in if num_documents_appear_in else 1))
        word_entry['TFIDF'] = word_entry['TFIDF LN Factor'] * word_entry['# of Appearances in Target Media']
    
    return word_list


"""
Saves output to Excel - list of words, in decreasing order of TFIDF.
"""
def save_to_excel(word_list, output_filename):
    df = pd.DataFrame(word_list)
    df = df.sort_values(by='TFIDF', ascending=False)
    df.to_excel(output_filename, index=False)


def main():
    known_words = load_known_words(KNOWN_WORDS_FILENAME)
    
    word_freq = calculate_target_media_word_frequency(INPUT_FILENAME)
    word_list = [{'word': word, '# of Appearances in Target Media': freq} for word, freq in word_freq.items() if word not in known_words]
    
    document_list = pull_reference_documents(DOCUMENTS_FOLDER)
    
    word_list = calculate_statistics(word_list, document_list)
    
    save_to_excel(word_list, OUTPUT_FILENAME)

if __name__ == "__main__":
    main()
